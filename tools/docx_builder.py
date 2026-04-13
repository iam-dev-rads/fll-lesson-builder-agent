import os
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configure logging
logger = logging.getLogger(__name__)

def _set_font(run, name="Calibri", size=12, bold=False):
    """
    Helper to set font properties on a run.
    """
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold

def _add_bordered_paragraph(doc, text, icon=""):
    """
    Adds a single-cell table to simulate a bordered paragraph/box.
    """
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # Add text with icon
    run = p.add_run(f"{icon} " if icon else "")
    _set_font(run, bold=True)
    run = p.add_run(text)
    
    return p

def _add_horizontal_line(doc):
    """
    Adds a horizontal line by setting a bottom border on a paragraph.
    """
    p = doc.add_paragraph()
    p_element = p._element
    p_pr = p_element.get_or_add_pPr()
    p_p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12') # thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    p_p_bdr.append(bottom)
    p_pr.append(p_p_bdr)

def build_coach_script(content: dict, output_folder: str) -> str:
    """
    Builds a Word document coach script based on the FLL lesson content JSON.
    Receives validated content dict and output folder path.
    Returns the absolute path to the saved .docx file.
    """
    doc = Document()
    
    # Set margins to 1 inch on all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. SECTION 1 — Header
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(content["lesson_title"])
    _set_font(run, name="Calibri", size=24, bold=True)
    
    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FLL Coach Script")
    _set_font(run, name="Calibri", size=16, bold=True)
    
    _add_horizontal_line(doc)

    # 2. SECTION 2 — Before You Begin
    doc.add_heading("Before You Begin", level=2)
    p = doc.add_paragraph(content["coach_script"]["intro"])
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)

    # 3. SECTION 3 — Opening The Lesson
    doc.add_heading("Opening Hook", level=2)
    p = doc.add_paragraph(content["opening_hook"])
    p.paragraph_format.line_spacing = 1.5
    
    _add_bordered_paragraph(
        doc, 
        "Coach Tip: Ask the kids this question and wait 30 seconds for answers before moving to the next slide.", 
        icon="💡"
    )

    # 4. SECTION 4 — Slide-by-Slide Script
    doc.add_page_break()
    main_heading = doc.add_heading("Slide-by-Slide Guide", level=1)
    main_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i, slide in enumerate(content["slides"]):
        # Sub-heading: Slide N: Title
        doc.add_heading(f"Slide {i+1}: {slide['title']}", level=3)
        
        # Script text
        try:
            script_text = content["coach_script"]["per_slide"][i]
        except IndexError:
            script_text = "[Script missing for this slide]"
        
        p = doc.add_paragraph()
        run = p.add_run(script_text)
        _set_font(run, size=12)
        p.paragraph_format.line_spacing = 1.5
        
        # Key points box
        _add_bordered_paragraph(doc, "Key Points:", icon="📌")
        for bullet in slide["bullet_points"]:
            bp = doc.add_paragraph(bullet, style='List Bullet')
            bp.paragraph_format.line_spacing = 1.5
            
        # Example callout
        _add_bordered_paragraph(doc, f"Example to share with kids: {slide['example']}", icon="🔍")
        doc.add_paragraph() # Spacer between slides

    # 5. SECTION 5 — Closing The Lesson
    doc.add_heading("Closing The Lesson", level=2)
    p = doc.add_paragraph(content["coach_script"]["wrap_up"])
    p.paragraph_format.line_spacing = 1.5
    
    exercise = content["closing_exercise"]
    _add_bordered_paragraph(doc, f"Exercise: {exercise['title']}", icon="✅")
    
    p = doc.add_paragraph()
    run = p.add_run("Instructions: ")
    run.bold = True
    p.add_run(exercise["instructions"])
    
    p = doc.add_paragraph()
    run = p.add_run("Expected outcome: ")
    run.bold = True
    p.add_run(exercise["expected_outcome"])

    # Force Calibri and 1.5 spacing globally for all Normal paragraphs
    for p in doc.paragraphs:
        if p.style.name == 'Normal':
            for run in p.runs:
                if not run.font.name: # Don't override if specifically set
                    run.font.name = 'Calibri'
            p.paragraph_format.line_spacing = 1.5

    # Save the document
    output_path = os.path.join(output_folder, "coach_script.docx")
    abs_path = os.path.abspath(output_path)
    doc.save(abs_path)
    
    logger.info(f"Coach script saved to: {abs_path}")
    return abs_path
